import json
import numpy as np
from sklearn.neighbors import KNeighborsClassifier
from sklearn.metrics import pairwise_distances
import tensorflow as tf
from flask import Flask, request, jsonify
import cv2
import numpy as np
import joblib
from flask_cors import CORS
from ultralytics import YOLO
import jwt
import psycopg2
from flask_bcrypt import Bcrypt
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import datetime

from mailpass import password

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})  # Enable CORS for all origins
bcrypt = Bcrypt(app)

global Email
global conn
app.config["SECRET_KEY"] = "jshefhbkj5456789654vgbdjgf"

conn = psycopg2.connect(
    host="ep-silent-cake-a1g9wkzh.ap-southeast-1.aws.neon.tech",
    database="snapmark",
    user="snapmark_owner",
    password='wGKHm9vzX1In'
)

def connectDB():
    global conn
    cur = conn.cursor()
    cur.execute('''CREATE TABLE IF NOT EXISTS USERS(
                id uuid DEFAULT uuid_generate_v4() PRIMARY KEY,
                email varchar(255) UNIQUE,
                name varchar(255),
                password varchar(255)
    );''')
    conn.commit()
    cur.close()

import cv2
from docx import Document
from docx.shared import Inches
import smtplib
from email.message import EmailMessage

from docx.shared import Cm
from PIL import Image


# Define a global variable to store predictions and face images
accumulated_predictions = []
accumulated_face_images = []


def send_email_with_attachment(file_path):
    global Email
    print("Sending email!!")
    # Email credentials
    sender_email = Email
    receiver_email = Email
    # receiver_email = "vivekraj@iiitdwd.ac.in"
    
    # Create a multipart message
    message = MIMEMultipart()
    message["From"] = sender_email
    message["To"] = receiver_email
    message["Subject"] = "Live Class Predictions Report"

    # Add body to email
    body = "Please find the attached Live Class Predictions Report."
    message.attach(MIMEText(body, "plain"))

    # Open the file in binary mode
    with open(file_path, "rb") as attachment:
        # Add file as application/octet-stream
        part = MIMEBase("application", "octet-stream")
        part.set_payload(attachment.read())

    # Encode file in ASCII characters to send by email    
    encoders.encode_base64(part)

    # Add header as key/value pair to attachment part
    part.add_header(
        "Content-Disposition",
        f"attachment; filename= {file_path}",
    )

    # Add attachment to message and convert message to string
    message.attach(part)
    text = message.as_string()

    # Log in to the email server
    server = smtplib.SMTP("smtp.gmail.com", 587)
    server.starttls()
    server.login(sender_email, password)

    # Send email
    server.sendmail(sender_email, receiver_email, text)
    server.quit()


def create_word_document(predictions, face_images):
    now = datetime.datetime.now()
    timestamp = now.strftime("%Y-%m-%d_%H-%M-%S")
    file_name = f"live_class_predictions_{timestamp}.docx"

    document = Document()
    
    # Add heading
    document.add_heading('Live Class Predictions', level=1)
    
    # Add table
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Face Image'
    hdr_cells[1].text = 'Predictions'
    
    # Add data to table
    for i, (prediction, face_image) in enumerate(zip(predictions, face_images)):
        # Add a new row
        row = table.add_row().cells

        # Add the face image to the first column
        img_path = f'face_image_{i}.png'
        cv2.imwrite(img_path, face_image)
        img = Image.open(img_path)
        width, height = img.size
        aspect_ratio = width / height
        img_width = Cm(2)  # Set width to 2 cm
        img_height = img_width / aspect_ratio
        row[0].paragraphs[0].add_run().add_picture(img_path, width=img_width, height=img_height)

        # Add predictions to the second column
        row[1].text = prediction

    # Save the document
    document.save(file_name)
    send_email_with_attachment(file_name)


    


with open("reversed_obj_ids.json", "r") as file:
    obj_id_to_usn = json.load(file)
# print(obj_id_to_usn)

class CustomKNN(KNeighborsClassifier):
    def __init__(self, n_neighbors=5, **kwargs):
        super().__init__(n_neighbors=n_neighbors, **kwargs)

    def predict(self, X):
        labels = super().predict(X)
        distances = pairwise_distances(X, self._fit_X, metric='euclidean', n_jobs=-1).min(axis=1)
        nearest_neighbor_index = np.argmin(pairwise_distances(X, self._fit_X, metric='euclidean', n_jobs=-1), axis=1)
        return labels, distances, nearest_neighbor_index

with open("reversed_obj_ids.json", "r") as file:
    obj_id_to_usn = json.load(file)

with open("obj.names", "r") as file:
    values = [int(line.strip()) for line in file]

custom_dict = {i: obj_id_to_usn[str(values[i])] for i in range(len(values))}

def load_tflite_model(file = "facenet.tflite"):
    interpreter = tf.lite.Interpreter(model_path=file)
    interpreter.allocate_tensors()
    return interpreter

def preprocess_image(image, image_size=(160, 160), kernel_size=(5, 5), sigma_x=0.5, sigma_y=0.5):
    resized_image = cv2.resize(image, image_size)
    resized_image = cv2.GaussianBlur(resized_image, kernel_size, sigma_x, sigma_y)
    resized_image = cv2.cvtColor(resized_image, cv2.COLOR_BGR2RGB)
    resized_image = resized_image.astype('float32')
    mean, std = resized_image.mean(), resized_image.std()
    resized_image = (resized_image - mean) / std
    return resized_image

def tflite_predict(face_model, image):
    input_details = face_model.get_input_details()
    output_details = face_model.get_output_details()
    input_shape = input_details[0]['shape']
    input_data = image.reshape(input_shape)
    face_model.set_tensor(input_details[0]['index'], input_data)
    face_model.invoke()
    output = []
    output_data = face_model.get_tensor(output_details[0]['index'])
    for i in output_data:
        output.append(i)
    return np.stack(output)

def get_tflite_facenet_embedding(image,tflite_model, image_size = (160,160), kernel_size=(3,3), sigma_x=0.5, sigma_y=0.5):
    preprocessed_image = preprocess_image(image=image)
    embedding = tflite_predict(face_model=tflite_model, image=preprocessed_image)
    curr_embedding = np.array(embedding)
    return curr_embedding

def get_models_prediction(models, model_thr, embedding,label_encoder,debug=0):
    knn_pred = ""
    rf_pred = ""
    lr_pred = ""
    svm_l_pred = ""
    svm_rbf_pred = ""
    svm_ploy_pred = ""
    svm_sig_pred = ""
    ann1_pred = ""
    ann2_pred = ""
    cnn1_pred = ""
    cnn2_pred = ""
    
    for i,pair in enumerate(model_thr):
        prediction = ""
        model = models[i]
        if pair[0]=="models-pkl/custom_knn_model.pkl":
            result = model.predict(embedding)
            pred = result[0][0]
            distance = result[1][0]
            prediction=""
            if distance>pair[1]: 
                prediction = obj_id_to_usn["100"]
            else: 
                prediction = obj_id_to_usn[str(pred)]

        if pair[0]=="models-pkl/random_forest_model.pkl":
            result = model.predict(embedding) [0]
            decision_score = model.predict_proba(embedding)
            prediction=obj_id_to_usn[str(result)]
            if np.max(decision_score[0])<pair[1]: 
                prediction="UNKNOWN"
            
        if pair[0]=="models-pkl/logistic_regression_model.pkl" or pair[0]=="models-pkl/svm_linear_model.pkl" or pair[0]=="models-pkl/svm_rbf_model.pkl" or pair[0]=="models-pkl/svm_poly_model.pkl" or pair[0]=="models-pkl/svm_sigmoid_model.pkl":
            result = model.predict(embedding) [0]
            decision_score = model.decision_function(embedding)
            prediction=obj_id_to_usn[str(result)]
            if np.max(decision_score[0])<pair[1]: 
                prediction="UNKNOWN"

        if pair[0] == "models-pkl/ann1_model.pkl" or pair[0] == "models-pkl/ann2_model.pkl" or pair[0] == "models-pkl/cnn1_model.pkl" or pair[0]=="models-pkl/cnn2_model.pkl":
            result = model.predict(embedding) [0]
            result_max_index = np.argmax(result)
            if(np.max(result)<pair[1]): 
                prediction="UNKNOWN"
            else:
                prediction=obj_id_to_usn[str(label_encoder.inverse_transform([result_max_index])[0])]

        if pair[0]=="models-pkl/custom_knn_model.pkl":    
            knn_pred = prediction
        if pair[0]=="models-pkl/random_forest_model.pkl":    
            rf_pred = prediction
        if pair[0]=="models-pkl/logistic_regression_model.pkl":    
            lr_pred = prediction
        if pair[0]=="models-pkl/svm_linear_model.pkl":    
            svm_l_pred = prediction
        if pair[0]=="models-pkl/svm_rbf_model.pkl":    
            svm_rbf_pred = prediction
        if pair[0]=="models-pkl/svm_poly_model.pkl":    
            svm_ploy_pred = prediction
        if pair[0]=="models-pkl/svm_sigmoid_model.pkl":    
            svm_sig_pred = prediction
        if pair[0]=="models-pkl/ann1_model.pkl":    
            ann1_pred = prediction
        if pair[0]=="models-pkl/ann2_model.pkl":    
            ann2_pred = prediction
        if pair[0]=="models-pkl/cnn1_model.pkl":    
            cnn1_pred = prediction
        if pair[0]=="models-pkl/cnn2_model.pkl":
            cnn2_pred = prediction

    return [knn_pred,rf_pred,lr_pred,svm_l_pred,svm_rbf_pred,svm_ploy_pred,svm_sig_pred,ann1_pred,ann2_pred,cnn1_pred,cnn2_pred]

def get_predictions(image_copy, models, model_thr, label_encoder, tflite_model):
    # print(image_copy)
    image = cv2.resize(image_copy,(160,160))

    embedding = get_tflite_facenet_embedding(image=image, tflite_model=tflite_model)

    knn_pred, rf_pred, lr_pred, svm_l_pred, svm_rbf_pred, svm_ploy_pred, svm_sig_pred, ann1_pred, ann2_pred, cnn1_pred, cnn2_pred = get_models_prediction(models=models, model_thr=model_thr, embedding=embedding, label_encoder=label_encoder)
    
    return [knn_pred, rf_pred, lr_pred, svm_l_pred, svm_rbf_pred, svm_ploy_pred, svm_sig_pred, ann1_pred, ann2_pred, cnn1_pred, cnn2_pred]

model_thr=  [
    ("models-pkl/custom_knn_model.pkl", 0.78),
    ("models-pkl/random_forest_model.pkl", 0.16),
    ("models-pkl/logistic_regression_model.pkl", 2.8),
    ("models-pkl/svm_linear_model.pkl", 31.35),
    ("models-pkl/svm_rbf_model.pkl", 32.3),
    ("models-pkl/svm_poly_model.pkl", 30),
    ("models-pkl/svm_sigmoid_model.pkl", 31.35),
    ("models-pkl/ann1_model.pkl", 0.35),
    ("models-pkl/ann2_model.pkl", 0.30),
    ("models-pkl/cnn1_model.pkl", 0.35),
    ("models-pkl/cnn2_model.pkl", 0.55)
]

model_knn = joblib.load("models-pkl/custom_knn_model.pkl")
model_rf = joblib.load("models-pkl/random_forest_model.pkl")
model_lr = joblib.load("models-pkl/logistic_regression_model.pkl")
model_svm_l = joblib.load("models-pkl/svm_linear_model.pkl")
model_svm_rbf = joblib.load("models-pkl/svm_rbf_model.pkl")
model_svm_poly = joblib.load("models-pkl/svm_poly_model.pkl")
model_svm_sig = joblib.load("models-pkl/svm_sigmoid_model.pkl")
model_ann1 = joblib.load("models-pkl/ann1_model.pkl")
model_ann2 = joblib.load("models-pkl/ann2_model.pkl")
model_cnn1 = joblib.load("models-pkl/cnn1_model.pkl")
model_cnn2 = joblib.load("models-pkl/cnn2_model.pkl")
models = [model_knn,
              model_rf,
              model_lr,
              model_svm_l,
              model_svm_rbf,
              model_svm_poly,
              model_svm_sig,
              model_ann1,
              model_ann2,
              model_cnn1,
              model_cnn2]
tflite_model = load_tflite_model()

label_encoder = "models-pkl/le.pkl"
label_encoder = joblib.load(label_encoder)
















def get_yolo_faces(image):
    model = YOLO('yolov8n-face.pt')
    # print("Yolov8 model loaded successfully")
    # print("image loaded")

    if image is not None:
        # print("here")
        # image = cv2.resize(image, (640, 640))

        # print("image resized")
        faces=[]
        results = model(image)
        # print("here1")
        for result in results:
            # print("show ")

            for box in result.boxes.xywh:
                        box = box.cpu().numpy().tolist()
                        box = np.around(box, decimals=4).tolist()
                        x,y,w,h=int(box[0]),int(box[1]),int(box[2]),int(box[3])
                        # face_image = image[y:y + h, x:x + w]
                        # face_image = cv2.resize(face_image, (200, 200))
                        # # predictions=model2(face_image)
                        # # print(predictions)
                        # cv2.imshow("Processed Face", face_image)
                        # cv2.waitKey(0)
                        faces.append([x,y,w,h])

        return faces
    else:
        print(f"Error: Unable to load image '{image}'.")
        return None



def get_preds(face_images):
    embeddings = []
    predictions = []
    for i, face_image in enumerate(face_images):
        # cv2.imshow("get_preds", face_image)
        # cv2.waitKey(0)
        preds = get_predictions(image_copy=face_image,label_encoder=label_encoder,model_thr=model_thr,models=models, tflite_model=tflite_model)
        # print("Predictions: ", preds)
        prediction_str = f"C-KNN: {preds[0]}\nRF: {preds[1]}\nLR: {preds[2]}\nSVM-L: {preds[3]}\nSVM-R: {preds[4]}\nSVM-P: {preds[5]}\nSVM-S: {preds[6]}\nANN-1: {preds[7]}\nANN-2: {preds[8]}\nCNN-1: {preds[9]}\nCNN-2: {preds[10]}"
        predictions.append(prediction_str)
        # predictions.append("Face-"+str(i)+": "+', '.join(preds))
    # print("Predictions: ", predictions)
    return predictions

def predict(image):
    # Predict labels for the image
    # faces = detect_faces(image)
    faces = get_yolo_faces(image=image)
    face_images = [image[y-h//2:y+h//2, x-w//2:x+w//2] for (x, y, w, h) in faces]
    predicted_labels = get_preds(face_images)
    accumulated_predictions.extend(predicted_labels)
    accumulated_face_images.extend(face_images)
    # create_word_document(predictions=predicted_labels,face_images=face_images)
    return predicted_labels

@app.route('/api', methods=['GET'])
def return_ascii():
    d = {}
    input_chr = str(request.args['query'])
    answer = str(ord(input_chr))
    d['output'] = answer
    d['result'] = "21bcs052"
    return jsonify(d)

@app.route('/', methods=["GET"])
def hello():
    d = {}
    d['pred'] = "test"
    return jsonify(d)

@app.route('/register', methods=["POST"])
def registration():
    data = request.json
    email = data.get('email')
    name = data.get('name')
    password = data.get('password')

    # Check if email already exists
    cur = conn.cursor()
    cur.execute('SELECT * FROM USERS WHERE email = %s', (email,))
    existing_user = cur.fetchone()
    cur.close()

    if existing_user:
        return jsonify({'message': 'Email already exists', 'status': False})

    # Hash the password
    hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')

    # Insert new user
    cur = conn.cursor()
    cur.execute( 
        '''INSERT INTO USERS
        (email, name, password) VALUES (%s, %s, %s)''', 
        (email, name, hashed_password)) 
    conn.commit()
    cur.close()

    response_data = {'message': 'Registration successful', 'status': True}
    return jsonify(response_data)

@app.route('/login', methods=["POST"])
def login():
    global Email
    data = request.json
    email = data.get('email')
    password = data.get('password')

    print(email,password)

    # Check if email already exists
    cur = conn.cursor()
    cur.execute('SELECT * FROM USERS WHERE email = %s', (email,))
    existing_user = cur.fetchone()
    cur.close()

    if not existing_user:
        return jsonify({'message': 'Account not exist.', "status": False})

    get_pass = existing_user[3]
    is_valid = bcrypt.check_password_hash(get_pass, password)

    if not is_valid:
        return jsonify({'message': 'Wrong Password', "status": False})

    # Generate JWT token
    token = jwt.encode(
        {"id": existing_user[0]},
        app.config["SECRET_KEY"],
        algorithm="HS256"
    )

    Email= email

    return jsonify({'message': 'Login Successfull', "status": True, 'token': token})

@app.route('/predict', methods=['POST'])
def predict_endpoint():
    try:
        # Check if image file is present in the request
        if 'image' not in request.files:
            print("error!!")
            return jsonify({'error': 'No image provided'}), 400
        
        # Read image file from request
        image_file = request.files['image']
        image_np = np.frombuffer(image_file.read(), np.uint8)
        image = cv2.imdecode(image_np, cv2.IMREAD_COLOR)

        # Predict labels for the image
        prediction = predict(image)

        return jsonify({'predicted_labels': prediction}), 200
    
    except Exception as e:
        print(e)
        return jsonify({'error': str(e)}), 500

@app.route('/save', methods=['POST'])
def save_predictions():
    global accumulated_predictions
    global accumulated_face_images
    
    # Create and save the Word document with accumulated predictions and face images
    create_word_document(predictions=accumulated_predictions, face_images=accumulated_face_images)
    
    # Clear the accumulated data for the next batch of predictions
    accumulated_predictions = []
    accumulated_face_images = []
    
    return jsonify({'message': 'Predictions saved successfully'}), 200
if __name__ == '__main__':
    app.run(host='0.0.0.0', debug=True, use_reloader = True)
